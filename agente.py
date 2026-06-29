"""
Agente Automatizado de Fichas Inmobiliarias
Extrae datos de Zonaprop y genera fichas en Word
"""

import os
import json
import time
import argparse
import logging
import requests
import json5
from openai import OpenAI
from typing import Dict, List, Set
import re
import random
import shutil
import tempfile
from scrapingbee import ScrapingBeeClient
from playwright.sync_api import sync_playwright
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from pathlib import Path
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).parent
load_dotenv(BASE_DIR / '.env')
FICHAS_DIR = BASE_DIR / 'fichas'
IMAGE_TEMPLATES_FILE = FICHAS_DIR / 'image_templates.json'
IMAGE_TEMPLATES: Set[str] = set()
COOKIES_DIR = BASE_DIR / 'cookies'
COOKIES_DIR.mkdir(exist_ok=True)
COOKIES_FILE = BASE_DIR / 'zonaprop_cookies.json'
USER_DATA_DIR = None
PROFILE_DIR = None
PROFILE_CLONE_PATH = None
HEADLESS_MODE = True
HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
    'Accept-Language': 'es-ES,es;q=0.9,en;q=0.8',
    'Accept-Encoding': 'gzip, deflate, br',
    'Connection': 'keep-alive',
    'Upgrade-Insecure-Requests': '1',
    'Sec-Fetch-Dest': 'document',
    'Sec-Fetch-Mode': 'navigate',
    'Sec-Fetch-Site': 'none',
    'Sec-Fetch-User': '?1',
    'Cache-Control': 'max-age=0',
}
PROPIEDADES_FILE = BASE_DIR / 'propiedades_link.txt'
PROCESADOS_FILE = BASE_DIR / 'procesados.txt'
COOKIES_FILE = BASE_DIR / 'zonaprop_cookies.json'

driver = None

OPENAI_MODEL = os.environ.get('OPENAI_MODEL', 'gpt-4o-mini')
REQUIRED_FIELDS = [
    'precio', 'expensas', 'ambientes', 'metros_totales',
    'metros_cubiertos', 'ubicacion', 'orientacion', 'antiguedad',
    'descripcion_completa'
]

SECURITY_KEYWORDS = ['Verificación de seguridad', 'security check', 'seguridad en curso']
AUTO_DESC_MARKER = "AUTO_GENERATED"
GEOCODE_ENABLED = os.environ.get('GEOCODE_ENABLED', '').lower() in ('1', 'true', 'yes')
GEOCODE_COUNTRY = os.environ.get('GEOCODE_COUNTRY', 'Argentina')
GEOCODE_CITY = os.environ.get('GEOCODE_CITY', 'Buenos Aires')
SCRAPINGBEE_API_KEY = os.environ.get('SCRAPINGBEE_API_KEY')
SCRAPINGBEE_ONLY = os.environ.get('SCRAPINGBEE_ONLY', '').lower() in ('1', 'true', 'yes')
SCRAPINGBEE_FALLBACK = os.environ.get('SCRAPINGBEE_FALLBACK', '').lower() in ('1', 'true', 'yes')
SCRAPINGBEE_PREMIUM = os.environ.get('SCRAPINGBEE_PREMIUM', 'true').lower() in ('1', 'true', 'yes')
SCRAPINGBEE_STEALTH = os.environ.get('SCRAPINGBEE_STEALTH', 'true').lower() in ('1', 'true', 'yes')
SCRAPINGBEE_WAIT = int(os.environ.get('SCRAPINGBEE_WAIT', '12000'))
SCRAPINGBEE_TIMEOUT = int(os.environ.get('SCRAPINGBEE_TIMEOUT', '120000'))
PLAYWRIGHT_FALLBACK = os.environ.get('PLAYWRIGHT_FALLBACK', 'true').lower() in ('1', 'true', 'yes')


def _pagina_verificacion(html: str) -> bool:
    lower = html.lower()
    return any(keyword.lower() in lower for keyword in SECURITY_KEYWORDS)


def _calcular_porcentaje_completado(datos: dict) -> tuple[int, float]:
    completados = sum(1 for campo in REQUIRED_FIELDS if datos.get(campo))
    total = len(REQUIRED_FIELDS)
    porcentaje = (completados / total) * 100 if total else 0
    return completados, porcentaje


def _registrar_metrica(direccion: str, completados: int, total: int, porcentaje: float):
    linea = f"{direccion}: {porcentaje:.1f}% ({completados}/{total})\n"
    meta_file = FICHAS_DIR / 'metricas_completitud.log'
    meta_file.parent.mkdir(parents=True, exist_ok=True)
    with open(meta_file, 'a', encoding='utf-8') as f:
        f.write(linea)


def cargar_templates_imagenes() -> None:
    if IMAGE_TEMPLATES_FILE.exists():
        try:
            with open(IMAGE_TEMPLATES_FILE, 'r', encoding='utf-8') as f:
                datos = json.load(f)
            for item in datos:
                IMAGE_TEMPLATES.add(item)
        except Exception:
            pass


def guardar_templates_imagenes() -> None:
    if not IMAGE_TEMPLATES:
        return
    with open(IMAGE_TEMPLATES_FILE, 'w', encoding='utf-8') as f:
        json.dump(list(IMAGE_TEMPLATES), f, indent=2)


def actualizar_templates(imagenes: List[str]) -> None:
    actualizado = False
    for url in imagenes:
        match = re.search(r'(https://imgar\.zonapropcdn\.com/avisos/)(\d+)(/.*)', url)
        if match:
            template = f"{match.group(1)}{{prop_id}}{match.group(3)}"
            if template not in IMAGE_TEMPLATES:
                IMAGE_TEMPLATES.add(template)
                actualizado = True
    if actualizado:
        guardar_templates_imagenes()


def _extraer_mapa(html: str, datos: dict):
    import base64
    lat_m = re.search(r"const\s+mapLatOf\s*=\s*['\"]([^'\"]+)['\"]", html)
    lng_m = re.search(r"const\s+mapLngOf\s*=\s*['\"]([^'\"]+)['\"]", html)
    url_m = re.search(r"const\s+urlMapOf\s*=\s*['\"]([^'\"]+)['\"]", html)
    if lat_m:
        try:
            datos['latitud'] = base64.b64decode(lat_m.group(1)).decode()
        except Exception:
            pass
    if lng_m:
        try:
            datos['longitud'] = base64.b64decode(lng_m.group(1)).decode()
        except Exception:
            pass
    if url_m:
        try:
            datos['url_mapa'] = base64.b64decode(url_m.group(1)).decode()
        except Exception:
            pass


def _extraer_id_propiedad(url: str) -> str | None:
    match = re.search(r'-(\d+)\.html', url)
    if match:
        return match.group(1)
    return None


def intentar_imagenes_por_patron(url: str) -> List[str]:
    prop_id = _extraer_id_propiedad(url)
    if not prop_id or not IMAGE_TEMPLATES:
        return []
    encontrados: List[str] = []
    for template in IMAGE_TEMPLATES:
        candidate = template.format(prop_id=prop_id)
        if candidate in encontrados:
            continue
        try:
            response = requests.head(candidate, headers=HEADERS, timeout=15)
            if response.status_code == 200:
                encontrados.append(candidate)
        except Exception:
            continue
    return encontrados


def _extraer_objeto_js(texto: str, identificador: str) -> str | None:
    idx = texto.find(identificador)
    if idx == -1:
        return None
    inicio = texto.find('{', idx)
    if inicio == -1:
        return None
    profundidad = 0
    for pos in range(inicio, len(texto)):
        char = texto[pos]
        if char == '{':
            profundidad += 1
        elif char == '}':
            profundidad -= 1
            if profundidad == 0:
                return texto[inicio:pos + 1]
    return None


def obtener_aviso_info(soup: BeautifulSoup):
    for script in soup.find_all('script'):
        contenido = script.string or ''
        if 'const avisoInfo' not in contenido:
            continue
        objeto = _extraer_objeto_js(contenido, 'const avisoInfo = ')
        if not objeto:
            continue
        try:
            cleaned = objeto.replace('urlMapOf', '""').replace('mapLatOf', '0').replace('mapLngOf', '0')
            return json5.loads(cleaned)
        except Exception:
            continue
    return None


def limpiar_texto_html(texto: str) -> str:
    if not texto:
        return ''
    return BeautifulSoup(texto, 'html.parser').get_text(separator='\n').strip()


def actualizar_datos_con_aviso(datos: dict, aviso: dict):
    def obtener_valor(code: str) -> str:
        caracteristica = aviso.get('mainFeatures', {}).get(code, {})
        return (caracteristica.get('value') or caracteristica.get('label') or '')

    precio = aviso.get('price')
    if not precio:
        precios = aviso.get('pricesData', [])
        if precios and precios[0].get('prices'):
            primera = precios[0]['prices'][0]
            precio = primera.get('formattedAmount') or primera.get('amount')
    if precio:
        datos['precio'] = str(precio).replace('$', '').strip()
    else:
        datos['precio'] = datos.get('precio', '')

    expensas = aviso.get('expenses')
    if expensas:
        try:
            expensas = int(expensas)
            datos['expensas'] = f"{expensas:,}".replace(',', '.')
        except Exception:
            datos['expensas'] = str(expensas)

    operacion = ''
    precios_info = aviso.get('pricesData', [])
    if precios_info:
        operacion = precios_info[0].get('operationType', {}).get('name', '')
    if not operacion:
        operacion = aviso.get('postingType', '')
    if operacion:
        datos['operacion'] = operacion.capitalize()

    direccion = aviso.get('address', {}).get('name')
    if direccion:
        datos['direccion'] = direccion
    barrio = aviso.get('location', {}).get('name')
    if barrio:
        datos['barrio'] = barrio

    datos['metros_totales'] = obtener_valor('CFT100')
    datos['metros_cubiertos'] = obtener_valor('CFT101')
    datos['ambientes'] = obtener_valor('CFT1')
    datos['banos'] = obtener_valor('CFT3')
    datos['dormitorios'] = obtener_valor('CFT2')
    antiguedad = obtener_valor('CFT5')
    if antiguedad:
        datos['antiguedad'] = f"{antiguedad} años" if not antiguedad.endswith('años') else antiguedad
    datos['ubicacion'] = obtener_valor('1000019')
    datos['orientacion'] = obtener_valor('1000029')

    descripcion = limpiar_texto_html(aviso.get('description', ''))
    if descripcion:
        datos['caracteristicas'] = descripcion
        datos['descripcion_completa'] = descripcion

    datos['titulo'] = aviso.get('postingTitle', datos.get('titulo', ''))
def _guardar_cookies(cookies: list[dict]):
    sanitized = []
    for cookie in cookies:
        clean = {k: v for k, v in cookie.items() if k not in {'sameSite', 'sameSitePolicy'}}
        sanitized.append(clean)
    timestamp = int(time.time())
    path = COOKIES_DIR / f"session_{timestamp}.json"
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(sanitized, f, indent=2)
    with open(COOKIES_FILE, 'w', encoding='utf-8') as f:
        json.dump(sanitized, f, indent=2)


def _cargar_cookies(driver):
    cookies_files = []
    if COOKIES_FILE.exists():
        cookies_files.append(COOKIES_FILE)
    cookies_files.extend(sorted(COOKIES_DIR.glob('session_*.json')))
    random.shuffle(cookies_files)
    cookies = []
    for file in cookies_files:
        try:
            with open(file, 'r', encoding='utf-8') as f:
                cookies.extend(json.load(f))
        except Exception:
            continue
    if not cookies:
        return
    driver.get('https://www.zonaprop.com.ar/')
    for cookie in cookies:
        try:
            driver.add_cookie(cookie)
        except Exception:
            continue
    driver.refresh()


def _cargar_cookies_str(max_cookies: int = 5) -> str:
    cookies_files = []
    if COOKIES_FILE.exists():
        cookies_files.append(COOKIES_FILE)
    cookies_files.extend(sorted(COOKIES_DIR.glob('session_*.json')))
    if not cookies_files:
        return ''
    try:
        all_cookies = []
        for f in cookies_files[:3]:
            try:
                all_cookies.extend(json.loads(Path(f).read_text(encoding='utf-8')))
            except Exception:
                continue
        selected = random.sample(all_cookies, min(max_cookies, len(all_cookies)))
        return '; '.join(f"{c['name']}={c['value']}" for c in selected if 'name' in c and 'value' in c)
    except Exception:
        return ''


def solicitar_login():
    global driver
    close_driver()
    init_driver(headless=False, user_data_dir=USER_DATA_DIR, profile_dir=PROFILE_DIR)
    logger.info("Resolvé el control y logueate en Zonaprop (captcha si aparece)")
    driver.get('https://www.zonaprop.com.ar/usuarios/login')
    input("Presioná Enter cuando Zonaprop ya no muestre 'Verificación de seguridad'.")
    cookies = driver.get_cookies()
    if cookies:
        _guardar_cookies(cookies)
        logger.info("Cookies guardadas para sesiones futuras")
    close_driver()
    init_driver(headless=True, user_data_dir=USER_DATA_DIR, profile_dir=PROFILE_DIR)


def _safe_copy_profile(src: Path, dest: Path) -> None:
    skip_dirs = {"Cache", "Code Cache", "GPUCache", "Crashpad", "GrShaderCache"}
    skip_files = {"SingletonLock", "SingletonCookie", "SingletonSocket", "LOCK", "lock"}
    for root, dirs, files in os.walk(src):
        dirs[:] = [d for d in dirs if d not in skip_dirs]
        rel = Path(root).relative_to(src)
        target_root = dest / rel
        target_root.mkdir(parents=True, exist_ok=True)
        for name in files:
            if name in skip_files:
                continue
            source_file = Path(root) / name
            target_file = target_root / name
            try:
                shutil.copy2(source_file, target_file)
            except Exception:
                continue


def _clone_profile(user_data_dir: str | None, profile_dir: str | None) -> str | None:
    global PROFILE_CLONE_PATH
    if not user_data_dir:
        return None
    if PROFILE_CLONE_PATH and PROFILE_CLONE_PATH.exists():
        return str(PROFILE_CLONE_PATH / (profile_dir or 'Default'))
    tmp = Path(tempfile.mkdtemp(prefix='zonaprop_profile_'))
    src_root = Path(user_data_dir)
    src_profile = src_root / (profile_dir or 'Default')
    dest_profile = tmp / (profile_dir or 'Default')
    try:
        _safe_copy_profile(src_profile, dest_profile)
        local_state = src_root / 'Local State'
        if local_state.exists():
            shutil.copy2(local_state, tmp / 'Local State')
        PROFILE_CLONE_PATH = tmp
        return str(dest_profile)
    except Exception as exc:
        logger.warning(f"No se pudo clonar el perfil: {exc}")
        shutil.rmtree(tmp, ignore_errors=True)
        return None


def init_driver(headless=None, user_data_dir=None, profile_dir=None):
    global driver, USER_DATA_DIR, PROFILE_DIR, PROFILE_CLONE_PATH
    if driver:
        close_driver()
    USER_DATA_DIR = user_data_dir
    PROFILE_DIR = profile_dir
    use_headless = HEADLESS_MODE if headless is None else headless
    options = Options()
    if use_headless:
        options.add_argument('--headless=new')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--disable-gpu')
    options.add_argument('--window-size=1920,1080')
    options.add_argument('--start-maximized')
    options.add_argument('--disable-blink-features=AutomationControlled')
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_experimental_option('useAutomationExtension', False)
    options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
    profile_clone = _clone_profile(USER_DATA_DIR, PROFILE_DIR)
    if profile_clone:
        clone_path = Path(profile_clone)
        options.add_argument(f"--user-data-dir={clone_path.parent}")
        if PROFILE_DIR:
            options.add_argument(f"--profile-directory={clone_path.name}")
    else:
        if USER_DATA_DIR:
            options.add_argument(f"--user-data-dir={USER_DATA_DIR}")
        if PROFILE_DIR:
            options.add_argument(f"--profile-directory={PROFILE_DIR}")
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=options)
    if use_headless:
        driver.execute_cdp_cmd('Page.addScriptToEvaluateOnNewDocument', {
            'source': '''
                Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
                Object.defineProperty(navigator, 'plugins', {get: () => [1, 2, 3]});
                Object.defineProperty(navigator, 'languages', {get: () => ['es-ES', 'es', 'en']});
                window.navigator.chrome = { runtime: {} };
            '''
        })
    _cargar_cookies(driver)
    logger.info("Navegador Chrome inicializado (modo sigiloso)")


def close_driver():
    global driver, PROFILE_CLONE_PATH
    if driver:
        driver.quit()
        driver = None
    if PROFILE_CLONE_PATH and PROFILE_CLONE_PATH.exists():
        shutil.rmtree(PROFILE_CLONE_PATH, ignore_errors=True)
        PROFILE_CLONE_PATH = None


def cargar_urls(archivo):
    if not archivo.exists():
        return []
    with open(archivo, 'r', encoding='utf-8') as f:
        return [line.strip() for line in f if line.strip()]


def cargar_procesados():
    return set(cargar_urls(PROCESADOS_FILE))


def guardar_procesado(url):
    with open(PROCESADOS_FILE, 'a', encoding='utf-8') as f:
        f.write(f"{url}\n")


def obtener_urls_nuevas():
    todos = cargar_urls(PROPIEDADES_FILE)
    procesados = cargar_procesados()
    return [url for url in todos if url not in procesados]


def crear_carpetas(direccion):
    carpeta = FICHAS_DIR / direccion
    imagenes_dir = carpeta / 'Imagenes'
    imagenes_dir.mkdir(parents=True, exist_ok=True)
    return carpeta


def extraer_datos_zonaprop(url):
    global driver
    
    import random
    
    logger.info(f"Extrayendo datos de: {url}")

    if SCRAPINGBEE_API_KEY:
        try:
            client = ScrapingBeeClient(api_key=SCRAPINGBEE_API_KEY)
            cookies_str = _cargar_cookies_str()
            params = {
                'render_js': 'true',
                'wait': SCRAPINGBEE_WAIT,
                'wait_for': 'h1, .posting-title, [data-qa="posting-title"], #description',
                'block_resources': 'false',
                'country_code': 'ar',
                'premium_proxy': str(SCRAPINGBEE_PREMIUM).lower(),
                'stealth_proxy': str(SCRAPINGBEE_STEALTH).lower(),
                'timeout': SCRAPINGBEE_TIMEOUT,
                'device': 'desktop',
            }
            if cookies_str:
                params['cookies'] = cookies_str
            response = client.get(url, params=params)
            if response.status_code == 200:
                html = response.text
                if _pagina_verificacion(html):
                    logger.warning("ScrapingBee devolvió pantalla de verificación; intentando fallback")
                else:
                    soup = BeautifulSoup(html, 'html.parser')
                    datos = {
                        'url': url,
                        'direccion': '',
                        'barrio': '',
                        'tipo': 'Departamento',
                        'operacion': '',
                        'precio': '',
                        'expensas': '',
                        'piso': '',
                        'ambientes': '',
                        'metros_totales': '',
                        'metros_cubiertos': '',
                        'ubicacion': '',
                        'orientacion': '',
                        'antiguedad': '',
                        'dormitorios': '',
                        'banos': '',
                        'apto_credito': '',
                        'ninos': '',
                        'mascotas': '',
                        'caracteristicas': '',
                        'imagenes': [],
                        'descripcion_completa': '',
                        'verificacion_activa': False,
                        'latitud': '',
                        'longitud': '',
                        'url_mapa': '',
                        'page_html': html
                    }
                    aviso_info = obtener_aviso_info(soup)
                    if aviso_info:
                        actualizar_datos_con_aviso(datos, aviso_info)
                    titulo = soup.find('h1')
                    if titulo:
                        datos['titulo'] = titulo.get_text(strip=True)
                    for img in soup.find_all('img'):
                        src = img.get('src') or img.get('data-src') or img.get('data-lazy') or img.get('data-srcset')
                        if src and 'zonaprop' in src and 'avisos' in src:
                            if src not in datos['imagenes']:
                                datos['imagenes'].append(src)
                    for meta in soup.find_all('meta', property='og:image'):
                        src = meta.get('content')
                        if src and 'avisos' in src:
                            datos['imagenes'].append(src)
                    _extraer_mapa(html, datos)
                    return datos
        except Exception as e:
            logger.warning(f"ScrapingBee falló, intentando fallback: {e}")

        if SCRAPINGBEE_ONLY and not SCRAPINGBEE_FALLBACK:
            logger.error("SCRAPINGBEE_ONLY activo y ScrapingBee falló. Abortando extracción.")
            return None

    if PLAYWRIGHT_FALLBACK:
        try:
            with sync_playwright() as p:
                browser = p.firefox.launch(headless=True)
                page = browser.new_page()
                page.goto(url, wait_until='networkidle', timeout=60000)
                try:
                    page.get_by_role("button", name="Leer descripción completa").click(timeout=5000)
                except Exception:
                    pass
                html = page.content()
                browser.close()
            if not _pagina_verificacion(html):
                soup = BeautifulSoup(html, 'html.parser')
                datos = {
                    'url': url,
                    'direccion': '',
                    'barrio': '',
                    'tipo': 'Departamento',
                    'operacion': '',
                    'precio': '',
                    'expensas': '',
                    'piso': '',
                    'ambientes': '',
                    'metros_totales': '',
                    'metros_cubiertos': '',
                    'ubicacion': '',
                    'orientacion': '',
                    'antiguedad': '',
                    'dormitorios': '',
                    'banos': '',
                    'apto_credito': '',
                    'ninos': '',
                    'mascotas': '',
                    'caracteristicas': '',
                    'imagenes': [],
                    'descripcion_completa': '',
                    'verificacion_activa': False,
                    'latitud': '',
                    'longitud': '',
                    'url_mapa': '',
                    'page_html': html
                }
                aviso_info = obtener_aviso_info(soup)
                if aviso_info:
                    actualizar_datos_con_aviso(datos, aviso_info)
                titulo = soup.find('h1')
                if titulo:
                    datos['titulo'] = titulo.get_text(strip=True)
                for img in soup.find_all('img'):
                    src = img.get('src') or img.get('data-src') or img.get('data-lazy') or img.get('data-srcset')
                    if src and 'zonaprop' in src and 'avisos' in src:
                        if src not in datos['imagenes']:
                            datos['imagenes'].append(src)
                for meta in soup.find_all('meta', property='og:image'):
                    src = meta.get('content')
                    if src and 'avisos' in src:
                        datos['imagenes'].append(src)
                _extraer_mapa(html, datos)
                return datos
            logger.warning("Playwright devolvió verificación; usando Selenium como último recurso")
        except Exception as e:
            logger.warning(f"Playwright falló, usando Selenium: {e}")

    if driver is None:
        init_driver()
    
    max_attempts = 3
    verificacion = False
    html = ''

    for intento in range(max_attempts):
        try:
            driver.get(url)
            time.sleep(3)
            time.sleep(random.uniform(2, 4))

            WebDriverWait(driver, 15).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )

            time.sleep(random.uniform(1, 2))

            try:
                descripcion_btn = WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, "//button[contains(normalize-space(.), 'Leer descripción completa')]"))
                )
                driver.execute_script("arguments[0].click();", descripcion_btn)
                time.sleep(0.5)
            except Exception:
                pass

            html = driver.page_source
            if _pagina_verificacion(html):
                verificacion = True
                logger.warning("Pantalla de verificación detectada; reintentando...")
                if intento < max_attempts - 1:
                    sleep_time = random.uniform(30, 60)
                    time.sleep(sleep_time)
                    driver.delete_all_cookies()
                    _cargar_cookies(driver)
                    driver.refresh()
                    continue
            else:
                verificacion = False
            break
        except Exception as e:
            logger.warning(f"Intento {intento + 1} fallido: {e}")
            if intento < max_attempts - 1:
                time.sleep(5)
                continue
            logger.error(f"Error al obtener la página después de {max_attempts} intentos: {e}")
            return None

    if not html:
        logger.error("No se pudo descargar el HTML")
        return None
    
    soup = BeautifulSoup(html, 'html.parser')
    
    datos = {
        'url': url,
        'direccion': '',
        'barrio': '',
        'tipo': 'Departamento',
        'operacion': '',
        'precio': '',
        'expensas': '',
        'piso': '',
        'ambientes': '',
        'metros_totales': '',
        'metros_cubiertos': '',
        'ubicacion': '',
        'orientacion': '',
        'antiguedad': '',
        'dormitorios': '',
        'banos': '',
        'apto_credito': '',
        'ninos': '',
        'mascotas': '',
        'caracteristicas': '',
        'imagenes': [],
        'descripcion_completa': '',
        'verificacion_activa': verificacion,
        'latitud': '',
        'longitud': '',
        'url_mapa': '',
        'page_html': html
    }

    _extraer_mapa(html, datos)

    aviso_info = obtener_aviso_info(soup)
    if aviso_info:
        actualizar_datos_con_aviso(datos, aviso_info)
    
    try:
        titulo = soup.find('h1')
        if titulo:
            datos['titulo'] = titulo.get_text(strip=True)
    except:
        pass
    
    if not datos.get('operacion'):
        precio_elem = soup.find(string=lambda t: t and ('alquiler' in t.lower() or 'venta' in t.lower()))
        if precio_elem:
            parent = precio_elem.find_parent()
            if parent:
                datos['operacion'] = precio_elem.strip().split()[0] if precio_elem.strip() else ''
    
    for item in soup.find_all(['span', 'div', 'li']):
        text = item.get_text(strip=True)
        
        if ('m² tot' in text or 'mts2' in text.lower()) and not datos.get('metros_totales'):
            partes = text.split()
            for i, p in enumerate(partes):
                if 'tot' in p.lower() or 'mts' in p.lower():
                    try:
                        datos['metros_totales'] = partes[i-1] if i > 0 else ''
                    except:
                        pass
        
        if 'amb' in text.lower() and any(c.isdigit() for c in text) and not datos.get('ambientes'):
            for c in text:
                if c.isdigit():
                    datos['ambientes'] = c
                    break
        
        if ('dorm' in text.lower() or 'dorm.' in text.lower()) and not datos.get('dormitorios'):
            for c in text:
                if c.isdigit():
                    datos['dormitorios'] = c
                    break
        
        if 'baño' in text.lower() and not datos.get('banos'):
            for c in text:
                if c.isdigit():
                    datos['banos'] = c
                    break
        
        if 'año' in text.lower() and any(c.isdigit() for c in text) and not datos.get('antiguedad'):
            for c in text:
                if c.isdigit():
                    datos['antiguedad'] = c + ' años'
                    break
    
    descripcion = soup.find('div', {'id': 'description'})
    if not descripcion:
        descripcion = soup.find('div', class_='description')
    if not descripcion:
        descripcion = soup.find('section', class_='description')
    
    if descripcion:
        datos['descripcion_completa'] = descripcion.get_text(strip=True)
        datos['caracteristicas'] = descripcion.get_text(strip=True)
    
    for img in soup.find_all('img'):
        src = img.get('src') or img.get('data-src') or img.get('data-lazy') or img.get('data-srcset')
        if src and 'zonaprop' in src and 'avisos' in src:
            if src not in datos['imagenes']:
                datos['imagenes'].append(src)
    
    for meta in soup.find_all('meta', property='og:image'):
        src = meta.get('content')
        if src and 'avisos' in src and src not in datos['imagenes']:
            datos['imagenes'].append(src)
    
    scripts = soup.find_all('script')
    for script in scripts:
        if script.string and 'imageUrls' in script.string:
            import re
            urls = re.findall(r'["\'](https://imgar[^\"\']+)["\']', script.string)
            for url in urls:
                if url not in datos['imagenes']:
                    datos['imagenes'].append(url)

    if not datos.get('caracteristicas'):
        descripcion_corta = soup.find('div', id='longDescription')
        if descripcion_corta:
            texto = descripcion_corta.get_text(separator='\n').strip()
            datos['caracteristicas'] = texto
            datos['descripcion_completa'] = texto
    
    return datos


def procesar_datos(datos):
    if not datos:
        return None

    url = datos.get('url', '')
    direccion = datos.get('direccion', '').strip()
    barrio = datos.get('barrio', '').strip()
    
    if not direccion:
        import re
        match = re.search(r'([A-Za-z]+\s+\d+)', url)
        if match:
            direccion = match.group(1).title()
        else:
            direccion = 'Propiedad'
        
        match = re.search(r'-([a-z]+)-\d+-', url)
        if match:
            direccion_match = re.search(r'([A-Z][a-z]+\s*\d+)', url)
            if direccion_match:
                direccion = direccion_match.group(1).replace('-', ' ').title()
        
        palabras = url.split('/')[-1].replace('.html', '').split('-')
        direccion_limpia = []
        for palabra in palabras:
            if palabra.isdigit():
                if len(palabra) >= 3:
                    direccion_limpia.append(palabra)
            elif any(c.isdigit() for c in palabra):
                direccion_limpia.append(palabra)
        if direccion_limpia:
            direccion = ' '.join(direccion_limpia[:2]).title()
        
        parts = url.split('/')[-1].replace('.html', '').replace('-', ' ').split()
        for i, part in enumerate(parts):
            if part.isdigit() and len(part) >= 3:
                direccion = f"{parts[i-1] if i > 0 else 'Calle'} {part}" if i > 0 else part
                break

        if 'larrea' in url.lower():
            direccion = 'Larrea 910'
            barrio = barrio or 'Recoleta'
        elif 'zapata' in url.lower():
            direccion = 'Zapata 561'
            barrio = barrio or 'Colegiales'
        elif 'cabrera' in url.lower():
            direccion = 'Cabrera 3200'
            barrio = barrio or 'Palermo'
        elif 'austria' in url.lower():
            direccion = 'Austria 2200'
            barrio = barrio or 'Palermo'
        elif 'matheu' in url.lower():
            direccion = 'Matheu 900'
            barrio = barrio or 'Almagro'
        elif 'caballito' in url.lower():
            direccion = 'Caballito'
            barrio = barrio or 'Caballito'
        elif 'urquiza' in url.lower():
            direccion = 'Villa Urquiza'
            barrio = barrio or 'Villa Urquiza'
        elif 'palermo' in url.lower():
            direccion = 'Palermo'
            barrio = barrio or 'Palermo'

    if 'alquiler' in url.lower():
        datos['operacion'] = 'Alquiler'
    elif 'venta' in url.lower():
        datos['operacion'] = 'Venta'
    
    datos['direccion'] = direccion
    datos['barrio'] = barrio
    
    if not datos.get('ambientes'):
        if '2-amb' in url.lower() or '2-ambientes' in url.lower():
            datos['ambientes'] = '2'
        elif '3-amb' in url.lower() or '3-ambientes' in url.lower():
            datos['ambientes'] = '3'
        elif '1-amb' in url.lower() or 'monoambiente' in url.lower():
            datos['ambientes'] = '1'
    
    return datos


def guardar_datos_raw(datos, carpeta):
    raw_file = carpeta / 'datos_raw.txt'
    with open(raw_file, 'w', encoding='utf-8') as f:
        f.write(f"URL: {datos.get('url', '')}\n\n")
        for key, value in datos.items():
            if key in ('page_html',):
                continue
            if key == 'imagenes':
                f.write(f"{key}:\n")
                for img_url in value:
                    f.write(f"  - {img_url}\n")
                continue
            f.write(f"{key}: {value}\n")


def guardar_html(datos, carpeta):
    html = datos.get('page_html')
    if not html:
        return
    html_file = carpeta / 'page.html'
    with open(html_file, 'w', encoding='utf-8') as f:
        f.write(html)


def _obtener_snippet(html: str, max_chars: int = 4000) -> str:
    return html[:max_chars].replace('\n', ' ') if html else ''


def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    r_pr.append(u)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    r_pr.append(color)

    new_run.append(r_pr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def obtener_sugerencia_llm(missing: List[str], html_snippet: str, direccion: str, completados: int, total: int) -> str | None:
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        return None

    client = OpenAI(api_key=api_key)
    prompt = (
        "Actúa como un scraper experto. Estoy extrayendo datos de Zonaprop y faltan los "
        f"campos {', '.join(missing)} para {direccion}. Tengo {completados}/{total} campos completados."
        f" Te paso un fragmento del HTML: {html_snippet}"
        "¿Qué selector o ajuste recomiendas para capturar esos elementos con Selenium + BeautifulSoup?"
    )
    try:
        completion = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "Eres un asistente técnico que sugiere ajustes de scraping."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=400,
        )
        return completion.choices[0].message.content
    except Exception as e:
        logger.warning(f"No se pudo consultar OpenAI: {e}")
        return None


def diagnosticar(datos, carpeta):
    issues = []
    missing = [campo for campo in REQUIRED_FIELDS if not datos.get(campo)]
    completados, porcentaje = _calcular_porcentaje_completado(datos)
    total = len(REQUIRED_FIELDS)

    if datos.get('verificacion_activa'):
        issues.append("Se detectó la pantalla de verificación; el scraping quedó bloqueado temporalmente.")

    if missing:
        issues.append(f"Faltan campos detectados: {', '.join(missing)}")

    html_path = carpeta / 'page.html'
    html_text = ''
    if html_path.exists():
        try:
            html_text = html_path.read_text(encoding='utf-8')
        except Exception:
            html_text = ''
    else:
        issues.append("No se guardó el HTML para análisis offline.")

    if missing and html_text:
        snippet = _obtener_snippet(html_text)
        sugerencia = obtener_sugerencia_llm(missing, snippet, datos.get('direccion', ''), completados, total)
        if sugerencia:
            issues.append(f"Sugerencia LLM: {sugerencia}")

    desc_file = carpeta / 'descripcion.txt'
    if not datos.get('descripcion_completa') and desc_file.exists():
        try:
            contenido = desc_file.read_text(encoding='utf-8')
            if AUTO_DESC_MARKER in contenido:
                issues.append("Completar descripcion.txt manualmente (fallback)")
        except Exception:
            pass

    _registrar_metrica(datos.get('direccion', 'Propiedad'), completados, total, porcentaje)

    diagnostico_file = carpeta / 'diagnostico.txt'
    with open(diagnostico_file, 'w', encoding='utf-8') as f:
        f.write(f"Completitud: {porcentaje:.1f}% ({completados}/{total})\n")
        if not issues:
            f.write("Diagnóstico: OK. No se detectaron problemas automáticos.\n")
        else:
            for linea in issues:
                f.write(f"- {linea}\n")

    if issues:
        logger.warning(f"Diagnóstico emitido en {diagnostico_file}: {issues}")
def guardar_descripcion(datos, carpeta):
    desc_file = carpeta / 'descripcion.txt'

    if desc_file.exists():
        try:
            contenido = desc_file.read_text(encoding='utf-8')
            if AUTO_DESC_MARKER not in contenido:
                logger.info(f"Descripción manual detectada, no se sobreescribe: {desc_file}")
                return
        except Exception:
            pass
    
    direccion = datos.get('direccion', '')
    barrio = datos.get('barrio', '')
    operacion = datos.get('operacion', '')
    precio = datos.get('precio', '')
    expensas = datos.get('expensas', '')
    piso = datos.get('piso', '')
    ambientes = datos.get('ambientes', '')
    m_totales = datos.get('metros_totales', '')
    m_cubiertos = datos.get('metros_cubiertos', '')
    ubicacion = datos.get('ubicacion', '')
    orientacion = datos.get('orientacion', '')
    antiguedad = datos.get('antiguedad', '')
    ninos = datos.get('ninos', '')
    mascotas = datos.get('mascotas', '')
    apto_credito = datos.get('apto_credito', '')
    caracteristicas = datos.get('caracteristicas', '')
    
    with open(desc_file, 'w', encoding='utf-8') as f:
        f.write(f"{AUTO_DESC_MARKER}\n")
        f.write(f"{direccion} ({barrio})\n")
        f.write(f"DEPARTAMENTO DE {ambientes} AMBIENTES- {ubicacion.upper()}- {m_totales} mts2 TOTALES -\n\n")
        
        if operacion and precio:
            f.write(f"{operacion}: $ {precio}\n")
        if expensas:
            f.write(f"Expensas: $ {expensas}\n")
        if piso:
            f.write(f"Piso: \t{piso}\n")
        if ambientes:
            f.write(f"Ambientes: {ambientes}\n")
        if m_totales:
            f.write(f"Metros totales: {m_totales} mts2\n")
        if m_cubiertos:
            f.write(f"Metros cubiertos: {m_cubiertos} mts2\n")
        if ubicacion:
            f.write(f"Ubicación: {ubicacion}\n")
        if orientacion:
            f.write(f"Orientación: {orientacion}\n")
        if antiguedad:
            f.write(f"Antigüedad: {antiguedad}\n")
        if operacion:
            f.write(f"Operación: {operacion}\n")
        if ninos:
            f.write(f"Niños: {ninos}\n")
        if mascotas:
            f.write(f"Mascotas: {mascotas}\n")
        if apto_credito:
            f.write(f"Apto crédito: {apto_credito}\n")
        
        f.write("\nCaracterísticas:\n")
        f.write(f"{caracteristicas}\n")


def aplicar_descripcion_manual(datos, carpeta):
    desc_file = carpeta / 'descripcion.txt'
    if not desc_file.exists():
        return
    try:
        contenido = desc_file.read_text(encoding='utf-8')
    except Exception:
        return
    if AUTO_DESC_MARKER in contenido:
        return

    lineas = contenido.splitlines()
    caracteristicas = []
    start = False
    for linea in lineas:
        if start:
            caracteristicas.append(linea)
        if linea.strip().lower().startswith('características:') or linea.strip().lower().startswith('caracteristicas:'):
            start = True
    texto = "\n".join([l for l in caracteristicas if l.strip()])
    if not texto:
        texto = contenido.strip()
    if texto:
        datos['caracteristicas'] = texto
        datos['descripcion_completa'] = texto
        logger.info(f"Descripción manual aplicada desde {desc_file}")


def inferir_piso(datos: dict) -> None:
    if datos.get('piso'):
        return
    texto = "\n".join([
        datos.get('caracteristicas', ''),
        datos.get('descripcion_completa', '')
    ])
    if not texto:
        return
    match = re.search(r"\b(\d+)\s*[º°]?\s*piso\b", texto, re.IGNORECASE)
    if match:
        datos['piso'] = f"{match.group(1)}°"


def inferir_metros(datos: dict) -> None:
    texto = "\n".join([
        datos.get('caracteristicas', ''),
        datos.get('descripcion_completa', '')
    ])
    if not texto:
        return
    texto_l = texto.lower()
    if not datos.get('metros_totales'):
        match = re.search(r"(\d+)\s*(?:m2|mts2|m²)\s*tot", texto_l)
        if match:
            datos['metros_totales'] = match.group(1)
    if not datos.get('metros_cubiertos'):
        match = re.search(r"(\d+)\s*(?:m2|mts2|m²)\s*cub", texto_l)
        if match:
            datos['metros_cubiertos'] = match.group(1)
    match = re.search(r"(\d+)\s*(?:m2|mts2|m²)\s*totales?\s*cubiertos?", texto_l)
    if match:
        if not datos.get('metros_totales'):
            datos['metros_totales'] = match.group(1)
        if not datos.get('metros_cubiertos'):
            datos['metros_cubiertos'] = match.group(1)


def inferir_campos_desde_texto(datos: dict) -> None:
    texto = "\n".join([
        datos.get('caracteristicas', ''),
        datos.get('descripcion_completa', '')
    ])
    if not texto:
        return
    texto_l = texto.lower()

    if not datos.get('ambientes'):
        match = re.search(r"(\d+)\s*amb\b", texto_l)
        if match:
            datos['ambientes'] = match.group(1)

    if not datos.get('banos'):
        match = re.search(r"(\d+)\s*bañ", texto_l)
        if match:
            datos['banos'] = match.group(1)

    if not datos.get('antiguedad'):
        match = re.search(r"(\d+)\s*años", texto_l)
        if match:
            datos['antiguedad'] = f"{match.group(1)} años"

    if not datos.get('ubicacion'):
        if 'frente' in texto_l:
            datos['ubicacion'] = 'Frente'
        elif 'contrafrente' in texto_l:
            datos['ubicacion'] = 'Contrafrente'

    if not datos.get('orientacion'):
        match = re.search(r"orientacion\s*:?\s*([a-z]+)", texto_l)
        if match:
            datos['orientacion'] = match.group(1).upper()

    if not datos.get('mascotas'):
        if 'no se aceptan mascotas' in texto_l or 'no acepta mascotas' in texto_l:
            datos['mascotas'] = 'No'
        elif 'acepta mascotas' in texto_l:
            datos['mascotas'] = 'Si'


def inferir_barrio(datos: dict) -> None:
    if datos.get('barrio') or not GEOCODE_ENABLED:
        return
    direccion = datos.get('direccion', '').strip()
    if not direccion:
        return
    query = f"{direccion}, {GEOCODE_CITY}, {GEOCODE_COUNTRY}"
    try:
        resp = requests.get(
            "https://nominatim.openstreetmap.org/search",
            params={
                "q": query,
                "format": "json",
                "addressdetails": 1,
                "limit": 1
            },
            headers={"User-Agent": "inmobiliaria-agent"},
            timeout=20
        )
        resp.raise_for_status()
        data = resp.json()
        if not data:
            return
        address = data[0].get("address", {})
        barrio = address.get("suburb") or address.get("neighbourhood") or address.get("city_district")
        if barrio:
            datos['barrio'] = barrio.title()
    except Exception:
        return


def _cargar_datos_raw(carpeta) -> dict:
    raw_file = carpeta / 'datos_raw.txt'
    if not raw_file.exists():
        return {}
    datos = {}
    try:
        lines = raw_file.read_text(encoding='utf-8').splitlines()
    except Exception:
        return {}
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            datos[key.strip()] = value.strip()
    return datos


def _parsear_descripcion_manual(carpeta) -> dict | None:
    desc_file = carpeta / 'descripcion.txt'
    if not desc_file.exists():
        return None
    contenido = desc_file.read_text(encoding='utf-8')
    if AUTO_DESC_MARKER in contenido:
        return None
    datos = _cargar_datos_raw(carpeta)
    datos.setdefault('direccion', carpeta.name.replace('_', ' '))
    datos.setdefault('barrio', '')
    datos.setdefault('tipo', 'Inmueble')
    datos.setdefault('operacion', '')
    datos.setdefault('precio', '')
    datos.setdefault('expensas', '')
    datos.setdefault('piso', '')
    datos.setdefault('ambientes', '')
    datos.setdefault('metros_totales', '')
    datos.setdefault('metros_cubiertos', '')
    datos.setdefault('ubicacion', '')
    datos.setdefault('orientacion', '')
    datos.setdefault('antiguedad', '')
    datos.setdefault('ninos', '')
    datos.setdefault('mascotas', '')
    datos.setdefault('apto_credito', '')
    datos.setdefault('caracteristicas', '')
    datos.setdefault('descripcion_completa', '')
    aplicar_descripcion_manual(datos, carpeta)
    inferir_piso(datos)
    return datos


def regenerar_word_desde_descripcion():
    if not FICHAS_DIR.exists():
        logger.info("No existe la carpeta fichas/")
        return
    for carpeta in FICHAS_DIR.iterdir():
        if not carpeta.is_dir():
            continue
        datos = _parsear_descripcion_manual(carpeta)
        if not datos:
            continue
        try:
            generar_word(datos, carpeta)
            logger.info(f"Word regenerado desde descripcion manual: {carpeta}")
        except PermissionError:
            logger.warning(f"No se pudo escribir el Word (archivo abierto): {carpeta}")
        except Exception as exc:
            logger.warning(f"Error regenerando Word en {carpeta}: {exc}")



def generar_word(datos, carpeta):
    direccion = datos.get('direccion', '')
    barrio = datos.get('barrio', '')
    operacion = datos.get('operacion', '')
    precio = datos.get('precio', '')
    expensas = datos.get('expensas', '')
    piso = datos.get('piso', '')
    ambientes = datos.get('ambientes', '')
    m_totales = datos.get('metros_totales', '')
    m_cubiertos = datos.get('metros_cubiertos', '')
    ubicacion = datos.get('ubicacion', '')
    orientacion = datos.get('orientacion', '')
    antiguedad = datos.get('antiguedad', '')
    ninos = datos.get('ninos', '')
    mascotas = datos.get('mascotas', '')
    apto_credito = datos.get('apto_credito', '')
    caracteristicas = datos.get('caracteristicas', '')
    tipo = datos.get('tipo', 'Inmueble')
    url = datos.get('url', '')

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    doc.add_heading(f"{direccion} ({barrio})", level=0)
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"{tipo.upper()} · {ambientes or '—'} ambientes · {ubicacion or 'Ubicación pendiente'}")
    subtitle_run.font.size = Pt(12)

    if url:
        p = doc.add_paragraph()
        p.add_run("URL: ")
        _add_hyperlink(p, url, url)

    imagenes_dir = carpeta / 'Imagenes'
    if imagenes_dir.exists():
        fotos = sorted(imagenes_dir.glob('*.jpg'))
        if fotos:
            try:
                doc.add_picture(str(fotos[0]), width=Inches(4.5))
            except Exception:
                doc.add_paragraph()
    else:
        doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    resumen_campos = [
        ("Operación", operacion or 'FALTA DATO'),
        ("Precio", f"$ {precio}" if precio else 'FALTA DATO'),
        ("Expensas", f"$ {expensas}" if expensas else 'FALTA DATO'),
        ("Piso", piso or 'FALTA DATO'),
        ("Ambientes", ambientes or 'FALTA DATO'),
        ("Metros totales", f"{m_totales} mts2" if m_totales else 'FALTA DATO'),
        ("Metros cubiertos", f"{m_cubiertos} mts2" if m_cubiertos else 'FALTA DATO'),
        ("Ubicación", ubicacion or 'FALTA DATO'),
        ("Orientación", orientacion or 'FALTA DATO'),
        ("Antigüedad", antiguedad or 'FALTA DATO'),
        ("Niños", ninos or 'FALTA DATO'),
        ("Mascotas", mascotas or 'FALTA DATO'),
        ("Apto crédito", apto_credito or 'FALTA DATO')
    ]

    for label, valor in resumen_campos:
        row = table.add_row().cells
        row[0].text = label
        run = row[1].paragraphs[0].add_run(valor)
        if valor == 'FALTA DATO':
            run.font.color.rgb = RGBColor(255, 0, 0)

    url_mapa = datos.get('url_mapa', '')
    if url_mapa:
        try:
            resp = requests.get(url_mapa, headers=HEADERS, timeout=15)
            if resp.status_code == 200 and len(resp.content) > 1000:
                mapa_path = carpeta / 'mapa.png'
                with open(mapa_path, 'wb') as f:
                    f.write(resp.content)
                doc.add_paragraph()
                doc.add_heading('Ubicación', level=2)
                doc.add_picture(str(mapa_path), width=Inches(4.5))
        except Exception:
            pass

    doc.add_paragraph()
    doc.add_heading('Características clave', level=2)
    if caracteristicas:
        for línea in caracteristicas.split('\n'):
            texto = línea.strip()
            if texto:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(texto)
    else:
        p = doc.add_paragraph()
        run = p.add_run('FALTA DATO')
        run.font.color.rgb = RGBColor(255, 0, 0)

    nota = doc.add_paragraph()
    nota.add_run('Ficha generada automáticamente. Verifique datos en la fuente original.').italic = True

    word_file = carpeta / f"{direccion}.docx"
    doc.save(word_file)
    logger.info(f"Word generado: {word_file}")


def _url_a_hd(url: str) -> str | None:
    # Format 1: ID ya separado en pares (como viene del HTML)
    # /avisos/1/00/58/48/15/79/720x532/2036102159.jpg
    m1 = re.search(r'/avisos/\d+/(\d+/\d+/\d+/\d+/\d+)/(\d+x\d+)/(\d+)\.jpg', url)
    if m1:
        pairs = m1.group(1)
        seq = m1.group(3)
        return f"https://imgar.zonapropcdn.com/avisos/resize/1/{pairs}/1200x1200/{seq}.jpg"
    # Format 2: ID sin separar
    # /avisos/1/00/58708332/1200x1200/2042378161.jpg
    m2 = re.search(r'/avisos/\d+/(\d+)/(\d+x\d+)/(\d+)\.jpg', url)
    if m2:
        prop_id = m2.group(1)
        size = m2.group(2)
        seq = m2.group(3)
        padded = prop_id.zfill(10)
        pairs = '/'.join(padded[i:i+2] for i in range(0, 10, 2))
        return f"https://imgar.zonapropcdn.com/avisos/resize/1/{pairs}/{size}/{seq}.jpg"
    return None


def _intentar_hd_adicionales(url_hd_base: str, carpeta: Path, max_fotos: int = 15):
    match = re.search(r'/resize/1/(\d+/\d+/\d+/\d+/\d+)/(\d+x\d+)/(\d+)\.jpg', url_hd_base)
    if not match:
        return
    pairs = match.group(1)
    size = match.group(2)
    base_seq = int(match.group(3))

    imagenes_dir = carpeta / 'Imagenes'
    existentes = len(list(imagenes_dir.glob('*.jpg')))
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Referer': 'https://www.zonaprop.com.ar/'
    }

    descargadas = 0
    for offset in range(1, 40):
        for seq in (base_seq + offset, base_seq - offset):
            if seq <= 0 or descargadas >= max_fotos:
                continue
            hd_url = f"https://imgar.zonapropcdn.com/avisos/resize/1/{pairs}/{size}/{seq}.jpg"
            try:
                time.sleep(random.uniform(0.3, 0.8))
                resp = requests.get(hd_url, headers=headers, timeout=15)
                if resp.status_code == 200 and len(resp.content) > 51200:
                    idx = existentes + descargadas + 1
                    nombre = f"foto_{idx:02d}.jpg"
                    ruta = imagenes_dir / nombre
                    with open(ruta, 'wb') as f:
                        f.write(resp.content)
                    descargadas += 1
                    logger.info(f"HD adicional descargada: {nombre} (seq={seq})")
            except Exception:
                continue
        if descargadas >= max_fotos:
            break


def descargar_imagenes(datos, carpeta):
    import random
    imagenes_dir = carpeta / 'Imagenes'
    imagenes = datos.get('imagenes', [])
    
    headers_img = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Referer': 'https://www.zonaprop.com.ar/'
    }
    
    primera_hd = None
    for i, url in enumerate(imagenes, 1):
        try:
            hd_url = _url_a_hd(url)
            download_url = url
            if hd_url:
                try:
                    hd_resp = requests.get(hd_url, headers=headers_img, timeout=15)
                    if hd_resp.status_code == 200 and len(hd_resp.content) > 51200:
                        download_url = hd_url
                        if not primera_hd:
                            primera_hd = hd_url
                except Exception:
                    pass

            time.sleep(random.uniform(0.5, 1.5))
            response = requests.get(download_url, headers=headers_img, timeout=30)
            response.raise_for_status()
            
            ext = '.jpg'
            if 'png' in download_url.lower():
                ext = '.png'
            elif 'webp' in download_url.lower():
                ext = '.webp'
            
            nombre = f"foto_{i:02d}{ext}"
            ruta = imagenes_dir / nombre
            
            with open(ruta, 'wb') as f:
                f.write(response.content)
            
            logger.info(f"Imagen descargada: {nombre}")
        except Exception as e:
            logger.error(f"Error descargando imagen {url}: {e}")

    if not imagenes:
        logger.warning("No se encontraron imágenes en la página")
    elif primera_hd:
        _intentar_hd_adicionales(primera_hd, carpeta)


def _generar_mapa_folium(datos: dict, carpeta: Path):
    lat = datos.get('latitud', '')
    lng = datos.get('longitud', '')
    if not lat or not lng:
        return
    try:
        import folium
        lat_f = float(lat)
        lng_f = float(lng)
        m = folium.Map(location=[lat_f, lng_f], zoom_start=16)
        folium.Marker([lat_f, lng_f], popup=datos.get('direccion', '')).add_to(m)
        mapa_html = carpeta / 'mapa_interactivo.html'
        m.save(str(mapa_html))
        logger.info(f"Mapa Folium generado: {mapa_html}")
    except Exception as e:
        logger.warning(f"No se pudo generar mapa Folium: {e}")


def procesar_propiedad(url):
    logger.info(f"Procesando propiedad: {url}")
    
    datos = extraer_datos_zonaprop(url)
    if not datos:
        logger.error(f"No se pudieron extraer datos de {url}")
        return False
    
    datos = procesar_datos(datos)
    if not datos:
        logger.error(f"Error al procesar datos de {url}")
        return False

    direccion = datos.get('direccion', 'propiedad')
    direccion_folder = direccion.replace(' ', '_').replace('/', '_')

    carpeta = crear_carpetas(direccion_folder)

    guardar_datos_raw(datos, carpeta)
    logger.info(f"Datos guardados en: {carpeta}")

    guardar_html(datos, carpeta)
    diagnosticar(datos, carpeta)

    guardar_descripcion(datos, carpeta)
    logger.info(f"Descripción guardada en: {carpeta / 'descripcion.txt'}")

    # REDAC-001: reescribir descripción con IA
    try:
        from fichas.rewriter_descripciones import rewrite_property
        rewrite_property(carpeta.name)
        logger.info(f"Descripción reescrita con IA para {carpeta.name}")
    except ImportError:
        logger.debug("rewriter_descripciones no disponible, se usa descripción original")
    except Exception as e:
        logger.warning(f"Error al reescribir descripción: {e}")

    # REDAC-001: generar narración con ElevenLabs
    try:
        from fichas.rewriter_audios import process_property as tts_property
        tts_property(carpeta.name)
        logger.info(f"Audio generado con ElevenLabs para {carpeta.name}")
    except ImportError:
        logger.debug("rewriter_audios no disponible, sin audio")
    except Exception as e:
        logger.warning(f"Error al generar audio: {e}")

    aplicar_descripcion_manual(datos, carpeta)
    inferir_piso(datos)
    inferir_metros(datos)
    inferir_campos_desde_texto(datos)
    inferir_barrio(datos)

    generar_word(datos, carpeta)
    _generar_mapa_folium(datos, carpeta)

    descargar_imagenes(datos, carpeta)

    guardar_procesado(url)
    logger.info(f"Propiedad procesada y marcada como procesada: {url}")
    
    return True


def procesar_todas():
    urls_nuevas = obtener_urls_nuevas()
    
    if not urls_nuevas:
        logger.info("No hay propiedades nuevas para procesar")
        return
    
    logger.info(f"Se encontraron {len(urls_nuevas)} propiedades nuevas")
    
    for url in urls_nuevas:
        procesar_propiedad(url)
        pause = random.uniform(60, 90)
        logger.info(f"Descanso breve para evitar detección: {pause:.1f}s")
        time.sleep(pause)


def modo_watch(intervalo=3600):
    logger.info(f"Iniciando modo monitoreo (cada {intervalo} segundos)")
    
    while True:
        try:
            procesar_todas()
        except Exception as e:
            logger.error(f"Error en el ciclo de monitoreo: {e}")
        
        logger.info(f"Esperando {intervalo} segundos...")
        time.sleep(intervalo)


def main():
    parser = argparse.ArgumentParser(description='Agente de Fichas Inmobiliarias')
    parser.add_argument('--watch', action='store_true', help='Modo monitoreo continuo')
    parser.add_argument('--interval', type=int, default=3600, help='Intervalo en segundos (default: 3600)')
    parser.add_argument('--login', action='store_true', help='Abrir navegador para iniciar sesión manual en Zonaprop (guarda cookies)')
    parser.add_argument('--user-data-dir', help='Ruta del perfil de Chrome/Chromium ya autenticado')
    parser.add_argument('--profile-directory', help='Nombre del subdirectorio de perfil dentro del user-data-dir')
    parser.add_argument('--rebuild-manual', action='store_true', help='Regenerar Word desde descripcion.txt manual')
    parser.add_argument('--rewrite', action='store_true', help='Reescribir descripciones de TODAS las propiedades con IA')
    parser.add_argument('--rewrite-prop', type=str, help='Reescribir descripción de una propiedad específica (nombre de carpeta)')
    parser.add_argument('--tts', action='store_true', help='Generar audios (narración ElevenLabs) para TODAS las propiedades')
    parser.add_argument('--tts-prop', type=str, help='Generar audio para una propiedad específica (nombre de carpeta)')
    
    args = parser.parse_args()

    FICHAS_DIR.mkdir(exist_ok=True)
    cargar_templates_imagenes()

    global USER_DATA_DIR, PROFILE_DIR
    USER_DATA_DIR = args.user_data_dir
    PROFILE_DIR = args.profile_directory

    if (args.login or not any(COOKIES_DIR.glob('session_*.json'))) and not SCRAPINGBEE_ONLY:
        logger.info("No hay cookies aceptadas; abriendo navegador para login manual y resolver el gate.")
        solicitar_login()

    if args.rebuild_manual:
        regenerar_word_desde_descripcion()
        return

    if args.rewrite_prop:
        from fichas.rewriter_descripciones import rewrite_property
        rewrite_property(args.rewrite_prop)
        print(f"Reescritura completada para {args.rewrite_prop}. Ejecutá update_web.py para ver cambios.")
        return

    if args.rewrite:
        from fichas.rewriter_descripciones import process_all as rewrite_all
        rewrite_all()
        return

    if args.tts_prop:
        from fichas.rewriter_audios import process_property
        process_property(args.tts_prop)
        print(f"Audio generado para {args.tts_prop}. Ejecutá update_web.py para sincronizar.")
        return

    if args.tts:
        from fichas.rewriter_audios import process_all as tts_all
        tts_all()
        return

    try:
        if args.watch:
            modo_watch(args.interval)
        else:
            procesar_todas()
    finally:
        close_driver()


if __name__ == '__main__':
    main()
